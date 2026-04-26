import os
import json
import subprocess
import tempfile
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Any, Union

from .logger import get_logger

logger = get_logger(__name__)


class MineruParser:
    __slots__ = ()

    @staticmethod
    def _run_mineru_command(
        input_path: Union[str, Path],
        output_dir: Union[str, Path],
        method: str = "auto",
        lang: Optional[str] = None,
        backend: str = "pipeline",
        start_page: Optional[int] = None,
        end_page: Optional[int] = None,
        formula: bool = True,
        table: bool = True,
        device: Optional[str] = None,
        source: str = "huggingface",
        vlm_url: Optional[str] = None,
    ) -> None:
        cmd = [
            "mineru",
            "-p", str(input_path),
            "-o", str(output_dir),
            "-m", method,
            "-b", backend,
        ]
        if source:
            cmd.extend(["--source", source])
        if lang:
            cmd.extend(["-l", lang])
        if start_page is not None:
            cmd.extend(["-s", str(start_page)])
        if end_page is not None:
            cmd.extend(["-e", str(end_page)])
        if not formula:
            cmd.extend(["-f", "false"])
        if not table:
            cmd.extend(["-t", "false"])
        if device:
            cmd.extend(["-d", device])
        if vlm_url:
            cmd.extend(["-u", vlm_url])

        try:
            env = os.environ.copy()
            env.setdefault("HF_HUB_OFFLINE", "1")
            env.setdefault("TRANSFORMERS_OFFLINE", "1")
            logger.info(f"[MinerU] Starting command: {' '.join(cmd)}")
            result = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                check=True,
                encoding="utf-8",
                errors="ignore",
                env=env,
                timeout=120,
            )
            logger.info(f"[MinerU] Command succeeded: stdout={result.stdout[:300] if result.stdout else '(empty)'}")
        except subprocess.TimeoutExpired as e:
            logger.error(
                f"[MinerU] TIMEOUT after 120s! Command: {' '.join(cmd)}"
            )
            if e.stdout:
                logger.error(f"[MinerU] Timeout stdout: {e.stdout[:500]}")
            if e.stderr:
                logger.error(f"[MinerU] Timeout stderr: {e.stderr[:500]}")
            raise RuntimeError(f"MinerU command timed out after 120s")
        except subprocess.CalledProcessError as e:
            logger.error(
                f"[MinerU] FAILED (exit {e.returncode})! Command: {' '.join(cmd)}"
            )
            logger.error(f"[MinerU] stderr: {e.stderr[:1000] if e.stderr else '(empty)'}")
            logger.error(f"[MinerU] stdout: {e.stdout[:1000] if e.stdout else '(empty)'}")
            raise
        except FileNotFoundError:
            raise RuntimeError(
                "mineru command not found. Please install MinerU 2.0:\n"
                "pip install -U 'mineru[core]'"
            )

    @staticmethod
    def _read_output_files(
        output_dir: Path, file_stem: str, method: str = "auto"
    ) -> Tuple[List[Dict[str, Any]], str]:
        md_file = output_dir / f"{file_stem}.md"
        json_file = output_dir / f"{file_stem}_content_list.json"

        subdir = output_dir / file_stem
        if subdir.exists():
            md_file = subdir / method / f"{file_stem}.md"
            json_file = subdir / method / f"{file_stem}_content_list.json"

        md_content = ""
        if md_file.exists():
            try:
                with open(md_file, "r", encoding="utf-8") as f:
                    md_content = f.read()
            except Exception as e:
                logger.warning(f"Could not read markdown file {md_file}: {e}")

        content_list = []
        if json_file.exists():
            try:
                with open(json_file, "r", encoding="utf-8") as f:
                    content_list = json.load(f)
            except Exception as e:
                logger.warning(f"Could not read JSON file {json_file}: {e}")

        return content_list, md_content

    @staticmethod
    def parse_pdf(
        pdf_path: Union[str, Path],
        output_dir: Optional[str] = None,
        method: str = "auto",
        lang: Optional[str] = None,
        **kwargs,
    ) -> Tuple[List[Dict[str, Any]], str]:
        try:
            pdf_path = Path(pdf_path)
            if not pdf_path.exists():
                raise FileNotFoundError(f"PDF file not found: {pdf_path}")

            name_without_suff = pdf_path.stem

            if output_dir:
                base_output_dir = Path(output_dir)
            else:
                base_output_dir = pdf_path.parent / "mineru_output"

            base_output_dir.mkdir(parents=True, exist_ok=True)

            MineruParser._run_mineru_command(
                input_path=pdf_path,
                output_dir=base_output_dir,
                method=method,
                lang=lang,
                **kwargs,
            )

            backend = kwargs.get("backend", "")
            if backend.startswith("vlm-"):
                method = "vlm"

            content_list, md_content = MineruParser._read_output_files(
                base_output_dir, name_without_suff, method=method
            )

            return content_list, md_content

        except Exception as e:
            logger.error(f"parse_pdf error: {e}")
            raise

    @staticmethod
    def parse_image(
        image_path: Union[str, Path],
        output_dir: Optional[str] = None,
        lang: Optional[str] = None,
        **kwargs,
    ) -> Tuple[List[Dict[str, Any]], str]:
        try:
            image_path = Path(image_path)
            if not image_path.exists():
                raise FileNotFoundError(f"Image file not found: {image_path}")

            mineru_supported = {".png", ".jpeg", ".jpg"}
            all_supported = {".png", ".jpeg", ".jpg", ".bmp", ".tiff", ".tif", ".gif", ".webp"}
            ext = image_path.suffix.lower()
            if ext not in all_supported:
                raise ValueError(f"Unsupported image format: {ext}")

            actual_image_path = image_path
            temp_converted_file = None

            if ext not in mineru_supported:
                try:
                    from PIL import Image
                except ImportError:
                    raise RuntimeError("PIL/Pillow required for image conversion. pip install Pillow")

                temp_dir = Path(tempfile.mkdtemp())
                temp_converted_file = temp_dir / f"{image_path.stem}_converted.png"
                try:
                    with Image.open(image_path) as img:
                        if img.mode in ("RGBA", "LA", "P"):
                            if img.mode == "P":
                                img = img.convert("RGBA")
                            background = Image.new("RGB", img.size, (255, 255, 255))
                            if img.mode == "RGBA":
                                background.paste(img, mask=img.split()[-1])
                            else:
                                background.paste(img)
                            img = background
                        elif img.mode not in ("RGB", "L"):
                            img = img.convert("RGB")
                        img.save(temp_converted_file, "PNG", optimize=True)
                        actual_image_path = temp_converted_file
                except Exception as e:
                    if temp_converted_file and temp_converted_file.exists():
                        temp_converted_file.unlink()
                    raise RuntimeError(f"Image conversion failed: {e}")

            name_without_suff = image_path.stem
            if output_dir:
                base_output_dir = Path(output_dir)
            else:
                base_output_dir = image_path.parent / "mineru_output"
            base_output_dir.mkdir(parents=True, exist_ok=True)

            try:
                MineruParser._run_mineru_command(
                    input_path=actual_image_path,
                    output_dir=base_output_dir,
                    method="ocr",
                    lang=lang,
                    **kwargs,
                )
                content_list, md_content = MineruParser._read_output_files(
                    base_output_dir, name_without_suff, method="ocr"
                )
                return content_list, md_content
            finally:
                if temp_converted_file and temp_converted_file.exists():
                    try:
                        temp_converted_file.unlink()
                        temp_converted_file.parent.rmdir()
                    except Exception:
                        pass
        except Exception as e:
            logger.error(f"parse_image error: {e}")
            raise

    @staticmethod
    def parse_office_doc(
        doc_path: Union[str, Path],
        output_dir: Optional[str] = None,
        lang: Optional[str] = None,
        **kwargs,
    ) -> Tuple[List[Dict[str, Any]], str]:
        try:
            doc_path = Path(doc_path)
            if not doc_path.exists():
                raise FileNotFoundError(f"Document not found: {doc_path}")

            supported = {".doc", ".docx", ".ppt", ".pptx", ".xls", ".xlsx"}
            if doc_path.suffix.lower() not in supported:
                raise ValueError(f"Unsupported office format: {doc_path.suffix}")

            import platform
            lo_timeout = 30 if platform.system() == "Windows" else 10
            candidate_cmds = []
            if platform.system() == "Windows":
                default_path = r"C:\Program Files\LibreOffice\program\soffice.exe"
                if os.path.isfile(default_path):
                    candidate_cmds.append(default_path)
            candidate_cmds.extend(["libreoffice", "soffice"])

            working_cmd = None
            for cmd in candidate_cmds:
                try:
                    result = subprocess.run(
                        [cmd, "--version"],
                        capture_output=True, check=True, timeout=lo_timeout,
                        encoding="utf-8", errors="ignore",
                    )
                    working_cmd = cmd
                    logger.info(f"LibreOffice detected: {cmd}")
                    break
                except (subprocess.CalledProcessError, FileNotFoundError, subprocess.TimeoutExpired):
                    continue

            if not working_cmd:
                raise RuntimeError(
                    "LibreOffice not found. Install it or convert to PDF manually."
                )

            with tempfile.TemporaryDirectory() as temp_dir:
                temp_path = Path(temp_dir)
                convert_cmd = [
                    working_cmd, "--headless", "--nologo", "--norestore",
                    "--convert-to", "pdf", "--outdir", str(temp_path), str(doc_path),
                ]
                result = subprocess.run(
                    convert_cmd, capture_output=True, text=True, timeout=60,
                    encoding="utf-8", errors="ignore",
                )
                if result.returncode != 0:
                    raise RuntimeError(f"LibreOffice conversion failed: {result.stderr}")

                pdf_files = list(temp_path.glob("*.pdf"))
                if not pdf_files:
                    raise RuntimeError(f"No PDF generated for {doc_path.name}")

                pdf_path = pdf_files[0]
                return MineruParser.parse_pdf(
                    pdf_path=pdf_path, output_dir=output_dir, lang=lang, **kwargs
                )
        except Exception as e:
            logger.error(f"parse_office_doc error: {e}")
            raise

    @staticmethod
    def parse_office_doc_python(
        doc_path: Union[str, Path],
        output_dir: Optional[str] = None,
        **kwargs,
    ) -> Tuple[List[Dict[str, Any]], str]:
        doc_path = Path(doc_path)
        if not doc_path.exists():
            raise FileNotFoundError(f"Document not found: {doc_path}")

        ext = doc_path.suffix.lower()
        md_parts = []
        content_list = []

        if ext == ".docx":
            from docx import Document as DocxDocument
            doc = DocxDocument(str(doc_path))
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                style_name = para.style.name if para.style else ""
                if style_name.startswith("Heading"):
                    try:
                        level = int(style_name.replace("Heading", "").strip())
                    except ValueError:
                        level = 1
                    md_parts.append(f"{'#' * level} {text}\n")
                else:
                    md_parts.append(f"{text}\n")
                content_list.append({"type": "text", "text": text})

            for table in doc.tables:
                md_parts.append("\n")
                table_data = []
                for i, row in enumerate(table.rows):
                    cells = [cell.text.strip() for cell in row.cells]
                    md_parts.append("| " + " | ".join(cells) + " |\n")
                    if i == 0:
                        md_parts.append("| " + " | ".join(["---"] * len(cells)) + " |\n")
                    table_data.append(cells)
                md_parts.append("\n")
                content_list.append({"type": "table", "table": table_data})

        elif ext == ".xlsx":
            from openpyxl import load_workbook
            wb = load_workbook(str(doc_path), data_only=True)
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                md_parts.append(f"## {sheet_name}\n\n")
                table_data = []
                for i, row in enumerate(ws.iter_rows(values_only=True)):
                    cells = [str(c) if c is not None else "" for c in row]
                    if not any(cells):
                        continue
                    md_parts.append("| " + " | ".join(cells) + " |\n")
                    if i == 0:
                        md_parts.append("| " + " | ".join(["---"] * len(cells)) + " |\n")
                    table_data.append(cells)
                md_parts.append("\n")
                content_list.append({"type": "table", "table": table_data, "sheet": sheet_name})
        else:
            raise ValueError(f"Python parsing not supported for: {ext}")

        md_content = "\n".join(md_parts)
        if output_dir:
            output_dir = Path(output_dir)
            output_dir.mkdir(parents=True, exist_ok=True)
            with open(output_dir / f"{doc_path.stem}_content_list.json", "w", encoding="utf-8") as f:
                json.dump(content_list, f, ensure_ascii=False, indent=2)

        return content_list, md_content
